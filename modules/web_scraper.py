"""
FILE: modules/web_scraper.py
PURPOSE: Scrape weather, news, and stock data
"""

import csv
import json
import logging
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path

import requests

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from config import (
    NEWS_RSS_URL, NEWS_SEARCH_URL,
    YAHOO_FINANCE_URL, SCRAPER_OUTPUT_DIR, PREFERENCES_FILE
)

logger = logging.getLogger("WebScraper")

WMO_DESCRIPTIONS = {
    0: "Clear sky", 1: "Mainly clear", 2: "Partly cloudy", 3: "Overcast",
    45: "Foggy", 48: "Icy fog", 51: "Light drizzle", 53: "Drizzle",
    55: "Heavy drizzle", 61: "Slight rain", 63: "Moderate rain",
    65: "Heavy rain", 71: "Slight snow", 73: "Moderate snow",
    75: "Heavy snow", 80: "Slight showers", 81: "Moderate showers",
    82: "Heavy showers", 95: "Thunderstorm", 96: "Thunderstorm with hail",
}

WMO_ICONS = {
    0: "☀️", 1: "🌤️", 2: "⛅", 3: "☁️",
    45: "🌫️", 48: "🌫️", 51: "🌦️", 53: "🌦️",
    55: "🌧️", 61: "🌧️", 63: "🌧️", 65: "🌧️",
    71: "🌨️", 73: "🌨️", 75: "❄️", 80: "🌦️",
    81: "🌧️", 82: "⛈️", 95: "⛈️", 96: "⛈️",
}


class WebScraper:
    """Scrape weather, news headlines, and stock prices."""

    def __init__(self):
        self.preferences = self._load_preferences()
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })

    def _load_preferences(self):
        try:
            if PREFERENCES_FILE.exists():
                with open(PREFERENCES_FILE) as f:
                    return json.load(f)
        except Exception:
            pass
        return {}

    def _save_preferences(self):
        try:
            with open(PREFERENCES_FILE, "w") as f:
                json.dump(self.preferences, f, indent=2)
        except Exception as e:
            logger.error(f"Preferences save failed: {e}")

    # ─── Weather ───────────────────────────────────────────

    def scrape_weather(self, city=None):
        """Fetch weather using Open-Meteo (free, no API key needed)."""
        city = city or self.preferences.get("details", {}).get("weather_city", "London")
        try:
            # Geocode city → lat/lon
            geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={requests.utils.quote(city)}&count=1"
            geo = self.session.get(geo_url, timeout=10).json()

            if not geo.get("results"):
                logger.error(f"City not found: {city}")
                return None

            loc = geo["results"][0]
            lat, lon = loc["latitude"], loc["longitude"]
            country = loc.get("country", "")

            # Fetch weather data
            url = (
                f"https://api.open-meteo.com/v1/forecast"
                f"?latitude={lat}&longitude={lon}"
                f"&current=temperature_2m,relative_humidity_2m,apparent_temperature,"
                f"weather_code,wind_speed_10m,uv_index"
                f"&daily=temperature_2m_max,temperature_2m_min,sunrise,sunset"
                f"&timezone=auto"
            )
            w = self.session.get(url, timeout=10).json()
            cur = w["current"]
            daily = w["daily"]

            code = cur.get("weather_code", 0)
            temp_c = cur.get("temperature_2m", "N/A")
            temp_f = round(temp_c * 9/5 + 32, 1) if isinstance(temp_c, (int, float)) else "N/A"

            result = {
                "city": f"{city}, {country}",
                "timestamp": datetime.now().isoformat(),
                "temp_c": temp_c,
                "temp_f": temp_f,
                "feels_like_c": cur.get("apparent_temperature", "N/A"),
                "humidity": cur.get("relative_humidity_2m", "N/A"),
                "description": WMO_DESCRIPTIONS.get(code, "Unknown"),
                "icon": WMO_ICONS.get(code, "🌡️"),
                "wind_kmph": cur.get("wind_speed_10m", "N/A"),
                "uv_index": cur.get("uv_index", "N/A"),
                "sunrise": daily.get("sunrise", [""])[0][11:16] if daily.get("sunrise") else "N/A",
                "sunset": daily.get("sunset", [""])[0][11:16] if daily.get("sunset") else "N/A",
                "max_temp_c": daily.get("temperature_2m_max", ["N/A"])[0],
                "min_temp_c": daily.get("temperature_2m_min", ["N/A"])[0],
            }

            logger.info(f"Weather scraped for {city}")
            return result

        except Exception as e:
            logger.error(f"Weather scrape failed ({city}): {e}")
            return None

    # ─── News ───────────────────────────────────────────────

    def scrape_news(self, topic=None, max_articles=10):
        """Fetch news headlines from Google News RSS."""
        topic = topic or self.preferences.get("details", {}).get("news_topic", "general")
        try:
            if topic and topic != "general":
                url = NEWS_SEARCH_URL.format(requests.utils.quote(topic))
            else:
                url = NEWS_RSS_URL

            resp = self.session.get(url, timeout=10)
            resp.raise_for_status()

            root = ET.fromstring(resp.content)
            items = root.findall(".//item")

            articles = []
            for item in items[:max_articles]:
                articles.append({
                    "title": item.findtext("title", ""),
                    "link": item.findtext("link", ""),
                    "published": item.findtext("pubDate", ""),
                    "source": item.findtext("source", ""),
                    "timestamp": datetime.now().isoformat(),
                })

            logger.info(f"News scraped: {len(articles)} articles")
            return articles

        except Exception as e:
            logger.error(f"News scrape failed: {e}")
            return []

    # ─── Stocks ─────────────────────────────────────────────

    def scrape_stocks(self, symbols=None):
        """Fetch stock prices from Yahoo Finance v8 API."""
        symbols = symbols or self.preferences.get("details", {}).get(
            "stock_symbols", ["AAPL", "GOOGL", "MSFT"]
        )

        results = []
        for symbol in symbols:
            try:
                url = f"https://query1.finance.yahoo.com/v8/finance/chart/{symbol}?interval=1d&range=2d"
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
                    "Accept": "application/json",
                }
                resp = requests.get(url, headers=headers, timeout=10)
                resp.raise_for_status()
                data = resp.json()

                meta = data["chart"]["result"][0]["meta"]
                price = meta.get("regularMarketPrice") or meta.get("chartPreviousClose", 0)
                prev = meta.get("previousClose") or meta.get("chartPreviousClose", price)
                change = round(price - prev, 2) if price and prev else 0
                change_pct = round((change / prev) * 100, 2) if prev else 0

                results.append({
                    "symbol": symbol,
                    "price": round(price, 2),
                    "prev_close": round(prev, 2),
                    "change": change,
                    "change_pct": change_pct,
                    "currency": meta.get("currency", "USD"),
                    "exchange": meta.get("exchangeName", ""),
                    "name": meta.get("shortName", symbol),
                    "timestamp": datetime.now().isoformat(),
                })
                logger.info(f"Stock scraped: {symbol} = {price}")

            except Exception as e:
                logger.error(f"Stock scrape failed ({symbol}): {e}")
                # Try fallback URL
                try:
                    url2 = f"https://query2.finance.yahoo.com/v8/finance/chart/{symbol}"
                    resp2 = requests.get(url2, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
                    data2 = resp2.json()
                    meta2 = data2["chart"]["result"][0]["meta"]
                    price2 = meta2.get("regularMarketPrice", 0)
                    results.append({
                        "symbol": symbol,
                        "price": round(price2, 2),
                        "prev_close": round(meta2.get("previousClose", price2), 2),
                        "change": 0, "change_pct": 0,
                        "currency": meta2.get("currency", "USD"),
                        "exchange": meta2.get("exchangeName", ""),
                        "name": meta2.get("shortName", symbol),
                        "timestamp": datetime.now().isoformat(),
                    })
                except Exception:
                    results.append({
                        "symbol": symbol, "price": "Error",
                        "error": str(e), "timestamp": datetime.now().isoformat(),
                    })

        return results

    # ─── Save Outputs ────────────────────────────────────────

    def save_to_csv(self, category, data):
        try:
            out_dir = SCRAPER_OUTPUT_DIR / category
            out_dir.mkdir(parents=True, exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            path = out_dir / f"{category}_{ts}.csv"
            rows = data if isinstance(data, list) else [data]
            if rows:
                # Get all unique keys from all rows to avoid fieldname mismatch
                fieldnames = []
                for row in rows:
                    for key in row.keys():
                        if key not in fieldnames:
                            fieldnames.append(key)
                
                with open(path, "w", newline="", encoding="utf-8") as f:
                    writer = csv.DictWriter(f, fieldnames=fieldnames)
                    writer.writeheader()
                    writer.writerows(rows)
            return str(path)
        except Exception as e:
            logger.error(f"CSV save failed: {e}")
        return None

    def save_to_docx(self, category, data):
        if not DOCX_AVAILABLE:
            return None
        try:
            out_dir = SCRAPER_OUTPUT_DIR / category
            out_dir.mkdir(parents=True, exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            path = out_dir / f"{category}_{ts}.docx"
            doc = Document()
            doc.add_heading(f"PyAutomate — {category.title()} Report", 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            rows = data if isinstance(data, list) else [data]
            for i, row in enumerate(rows, 1):
                doc.add_heading(f"Entry {i}", level=2)
                for k, v in row.items():
                    doc.add_paragraph(f"{k}: {v}")
            doc.save(str(path))
            return str(path)
        except Exception as e:
            logger.error(f"DOCX save failed: {e}")
        return None
