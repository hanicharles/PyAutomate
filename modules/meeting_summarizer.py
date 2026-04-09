"""
FILE: modules/meeting_summarizer.py
PURPOSE: Summarize meeting transcripts using HuggingFace Inference API
"""

import csv
import logging
import os
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import speech_recognition as sr
    SR_AVAILABLE = True
except Exception as e:
    SR_AVAILABLE = False
    print(f"DEBUG: speech_recognition import failed with: {e}")

try:
    from pydub import AudioSegment
    PYDUB_AVAILABLE = True
except ImportError:
    PYDUB_AVAILABLE = False

from huggingface_hub import InferenceClient

from config import MEETING_OUTPUT_DIR

logger = logging.getLogger("MeetingSummarizer")

HF_TOKEN = os.getenv("HF_TOKEN", "YOUR_HF_TOKEN")

SYSTEM_PROMPT = """You are an expert at extracting and summarizing transcripts into clear, structured notes with emojis.

First, detect the type of transcript:
- TYPE A: **Project Presentation / Feedback Session** — someone presenting a project and receiving feedback/corrections
- TYPE B: **Lecture / Class Session** — a teacher/professor teaching concepts, giving explanations, and assigning work

Then produce the summary in the matching format below.

═══════════════════════════════════════
FORMAT FOR TYPE A (Project Feedback):
═══════════════════════════════════════

📝 Summary – [Project Name] Presentation Feedback
📌 Project Overview
Topic: [Extract Topic]
Goal:
[List extracted goals as bullet points]

⚠️ Problems in Current (Manual) System
[List issues with the current/manual approach]

❌ Issues Identified in Your Project
1. [Category Name]
[List specific details as bullet points]
2. [Category Name]
[List specific details as bullet points]
(Add more numbered categories as needed)

💡 Suggestions for Improvement
[Provide actionable suggestions based on the feedback]

🎯 Key Takeaway
👉 [Single overall takeaway sentence]

═══════════════════════════════════════
FORMAT FOR TYPE B (Lecture / Class):
═══════════════════════════════════════

📘 Class Summary Notes ([Main Topics Covered])

🗂️ 1. [First Topic]
[Key points as bullet points]
[Sub-points indented]

🧠 2. [Second Topic]
[Key points as bullet points]

📊 3. [Comparisons / Differences if any]
🔹 [Category A]
[Details]
🔹 [Category B]
[Details]
👉 Key Difference:
[One-line summary of difference]

🧩 4. [Assignment / Task if mentioned]
🎯 Goal:
[What needs to be done]

⚙️ 5. Required Components
[Numbered list of components needed]

💡 6. Example Use Case
[Practical example mentioned in class]

🛠️ 7. Implementation Notes
[Technical guidance given by instructor]

⏰ 8. Deadline
[Any deadlines mentioned]

🧾 Key Takeaways
[Bullet-point summary of most important things from the session]

═══════════════════════════════════════

RULES:
- Use the appropriate format based on transcript type
- Use emojis as shown in the templates
- Extract ALL important information — do not skip topics
- Clean up filler words, repetitions, and transcription noise
- For lectures: organize by topic with numbered sections
- For feedback: group issues into clear categories
- Be thorough and detailed — capture every concept, assignment, and instruction"""


DETAILED_NOTES_PROMPT = """You are an expert educator who creates deep, exam-ready study notes from transcripts.

Given a transcript and its summary, generate DETAILED explanations of every topic mentioned.

RULES:
- For EACH topic/concept mentioned, provide:
  - What it is (definition)
  - Key features (with checkmarks)
  - Examples with code/JSON if applicable
  - Comparisons with tables where relevant
  - Architecture diagrams using text arrows (→, ↓)
- Use emojis: 📘 for topics, 🔹 for sub-points, ✅ for features, 👉 for key points
- Use numbered sections: 📘 1. Topic Name, 📘 2. Topic Name, etc.
- Include comparison tables using plain text alignment
- Include code snippets in ``` blocks where relevant
- Add a ✅ Final Key Points section at the end
- Be VERY thorough — explain as if teaching someone who missed the class
- Make it exam-ready with clear definitions and examples"""

CHATBOT_PROMPT = """You are a helpful teaching assistant. You have access to a class/meeting transcript and its summary.

Based ONLY on the information in the transcript and summary provided, answer the user's question clearly and concisely.

If the question is about a topic mentioned in the transcript, explain it in detail with examples.
If the question is not related to the transcript content, politely say that the topic was not covered in this session.

Use emojis sparingly for clarity. Be helpful and educational."""


class MeetingSummarizer:
    """Summarize meeting transcripts using HuggingFace LLM."""

    def __init__(self):
        self.transcript = ""
        self.summary = ""
        os.makedirs(MEETING_OUTPUT_DIR, exist_ok=True)
        self.client = InferenceClient(token=HF_TOKEN)

    def load_text_file(self, filepath):
        """Load transcript from a .txt file."""
        try:
            with open(filepath, encoding="utf-8") as f:
                self.transcript = f.read()
            logger.info(f"Loaded text file: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Load text failed: {e}")
            return False

    def set_transcript(self, text):
        """Set transcript directly from pasted text."""
        self.transcript = text.strip()
        return bool(self.transcript)

    def load_audio_file(self, filepath):
        """Transcribe audio file using SpeechRecognition."""
        global SR_AVAILABLE, sr

        if not SR_AVAILABLE:
            try:
                import speech_recognition as sr_retry
                sr = sr_retry
                SR_AVAILABLE = True
            except Exception as e:
                logger.error(f"SpeechRecognition not available: {e}")
                return False

        temp_wav = None
        try:
            filepath = Path(filepath)
            ext = filepath.suffix.lower()

            if ext != ".wav":
                if not PYDUB_AVAILABLE:
                    logger.error("pydub not installed, cannot convert audio")
                    return False
                audio = AudioSegment.from_file(str(filepath), format=ext.replace(".", ""))
                temp_wav = filepath.with_suffix(".temp.wav")
                audio.export(str(temp_wav), format="wav")
                process_path = str(temp_wav)
            else:
                process_path = str(filepath)

            recognizer = sr.Recognizer()
            with sr.AudioFile(process_path) as source:
                audio = recognizer.record(source)
            self.transcript = recognizer.recognize_google(audio)
            logger.info(f"Transcribed audio: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Audio transcription failed: {e}")
            return False
        finally:
            if temp_wav and temp_wav.exists():
                try:
                    os.remove(temp_wav)
                except:
                    pass

    def load_video_file(self, filepath):
        """Extract audio from video file and transcribe."""
        if not PYDUB_AVAILABLE:
            logger.error("pydub not installed, cannot extract audio from video")
            return False

        filepath = Path(filepath)
        audio_path = filepath.with_suffix(".temp.wav")
        try:
            video_audio = AudioSegment.from_file(str(filepath))
            video_audio.export(str(audio_path), format="wav")
            success = self.load_audio_file(audio_path)
            return success
        except Exception as e:
            logger.error(f"Video processing failed: {e}")
            return False
        finally:
            if audio_path.exists():
                try:
                    os.remove(audio_path)
                except:
                    pass

    def analyze_transcript(self):
        """Send transcript to HuggingFace and get formatted summary."""
        if not self.transcript:
            return None

        try:
            logger.info("Sending transcript to HuggingFace for summarization...")
            response = self.client.chat_completion(
                model="meta-llama/Meta-Llama-3-8B-Instruct",
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": f"Please summarize this transcript:\n\n{self.transcript}"}
                ],
                max_tokens=2000,
                temperature=0.3,
            )
            self.summary = response.choices[0].message.content
            # Clean up format header lines the model sometimes echoes back
            import re
            self.summary = re.sub(r'═{3,}.*\n?', '', self.summary)
            self.summary = re.sub(r'FORMAT FOR TYPE [AB].*\n?', '', self.summary)
            self.summary = self.summary.strip()
            logger.info("Transcript summarized successfully")
            return {"summary": self.summary}
        except Exception as e:
            logger.error(f"HuggingFace summarization failed: {e}")
            return None

    def generate_detailed_notes(self):
        """Generate deep, exam-ready detailed notes from the transcript and summary."""
        if not self.transcript:
            return None

        try:
            logger.info("Generating detailed notes...")
            context = f"TRANSCRIPT:\n{self.transcript}"
            if self.summary:
                context += f"\n\nSUMMARY:\n{self.summary}"

            response = self.client.chat_completion(
                model="meta-llama/Meta-Llama-3-8B-Instruct",
                messages=[
                    {"role": "system", "content": DETAILED_NOTES_PROMPT},
                    {"role": "user", "content": f"Generate detailed, exam-ready notes for all topics in this transcript:\n\n{context}"}
                ],
                max_tokens=3000,
                temperature=0.3,
            )
            result = response.choices[0].message.content
            import re
            result = re.sub(r'═{3,}.*\n?', '', result)
            result = result.strip()
            self.detailed_notes = result
            logger.info("Detailed notes generated successfully")
            return result
        except Exception as e:
            logger.error(f"Detailed notes generation failed: {e}")
            return None

    def save_detailed_notes(self):
        """Save detailed notes as TXT and DOCX."""
        if not self.detailed_notes:
            return None
        try:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            txt_path = MEETING_OUTPUT_DIR / f"detailed_notes_{ts}.txt"
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(self.detailed_notes)
            
            docx_path = None
            if DOCX_AVAILABLE:
                docx_path = MEETING_OUTPUT_DIR / f"detailed_notes_{ts}.docx"
                doc = Document()
                doc.add_heading("Detailed Study Notes", 0)
                doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")
                for line in self.detailed_notes.split("\n"):
                    doc.add_paragraph(line)
                doc.save(str(docx_path))
            
            return {"txt": str(txt_path), "docx": str(docx_path) if docx_path else None}
        except Exception as e:
            logger.error(f"Detailed notes save failed: {e}")
            return None

    def chat(self, question):
        """Answer a question based on the transcript and summary context."""
        if not self.transcript:
            return "No transcript loaded. Please paste or upload a transcript first."

        try:
            logger.info(f"Chatbot answering: {question[:50]}...")
            context = f"TRANSCRIPT:\n{self.transcript}"
            if self.summary:
                context += f"\n\nSUMMARY:\n{self.summary}"

            response = self.client.chat_completion(
                model="meta-llama/Meta-Llama-3-8B-Instruct",
                messages=[
                    {"role": "system", "content": CHATBOT_PROMPT + f"\n\nCONTEXT:\n{context}"},
                    {"role": "user", "content": question}
                ],
                max_tokens=1500,
                temperature=0.4,
            )
            answer = response.choices[0].message.content.strip()
            logger.info("Chatbot answered successfully")
            return answer
        except Exception as e:
            logger.error(f"Chatbot failed: {e}")
            return f"Sorry, I couldn't process your question. Error: {str(e)}"

    def save_txt(self):
        """Save summary as plain text."""
        if not self.summary:
            return None
        try:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            path = MEETING_OUTPUT_DIR / f"meeting_{ts}.txt"
            with open(path, "w", encoding="utf-8") as f:
                f.write(self.summary)
            logger.info(f"Saved TXT: {path}")
            return str(path)
        except Exception as e:
            logger.error(f"TXT save failed: {e}")
            return None

    def save_docx(self):
        """Save summary as Word document."""
        if not self.summary or not DOCX_AVAILABLE:
            return None
        try:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            path = MEETING_OUTPUT_DIR / f"meeting_{ts}.docx"
            doc = Document()
            doc.add_heading("Meeting Summary", 0)
            doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")
            for line in self.summary.split("\n"):
                doc.add_paragraph(line)
            doc.save(str(path))
            logger.info(f"Saved DOCX: {path}")
            return str(path)
        except Exception as e:
            logger.error(f"DOCX save failed: {e}")
            return None

    def save_csv(self):
        """Save summary as CSV."""
        if not self.summary:
            return None
        try:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            path = MEETING_OUTPUT_DIR / f"meeting_{ts}.csv"
            with open(path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["line"])
                for line in self.summary.split("\n"):
                    if line.strip():
                        writer.writerow([line])
            logger.info(f"Saved CSV: {path}")
            return str(path)
        except Exception as e:
            logger.error(f"CSV save failed: {e}")
            return None

    def run(self):
        """CLI entry point."""
        print("\n" + "═" * 55)
        print("  🎤 MEETING SUMMARIZER")
        print("═" * 55)

        print("""
  [1] 📄  Load text transcript
  [2] 📝  Paste/type transcript
  [3] 🔙  Back
        """)
        choice = input("  Choice: ").strip()

        if choice == "1":
            path = input("  File path: ").strip()
            self.load_text_file(path)
        elif choice == "2":
            print("  Enter transcript (type END on a new line to finish):")
            lines = []
            while True:
                line = input()
                if line.strip().upper() == "END":
                    break
                lines.append(line)
            self.transcript = "\n".join(lines)

        if self.transcript:
            result = self.analyze_transcript()
            if result:
                print(f"\n{result['summary']}")
                self.save_txt()
                self.save_docx()
                self.save_csv()
                print("\n  💾 Saved to meeting_output/")
